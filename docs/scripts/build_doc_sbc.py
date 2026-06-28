# -*- coding: utf-8 -*-
"""
Relatório técnico (formato de artigo SBC) descrevendo o desenvolvimento do
Cervix Visual AI: proposta, desenvolvimento e avaliação.
Fonte Times New Roman 12pt, página A4, coluna única, margens SBC.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.dirname(os.path.abspath(__file__))
BLACK = RGBColor(0, 0, 0)
HDR_FILL = "D9D9D9"   # cinza claro para cabeçalho de tabela (sóbrio, estilo SBC)
ALT_FILL = "F2F2F2"
FONT = "Times New Roman"

doc = Document()

# ---------------------------------------------------------------- estilos SBC
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.0

def _style(name, size, bold=True):
    st = doc.styles[name]
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.name = FONT
    st.font.color.rgb = BLACK
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.0
    try:
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    except Exception:
        pass

_style("Heading 1", 13)
_style("Heading 2", 12)
_style("Heading 3", 12)

# margens SBC: A4, superior 3,5 cm; esquerda/direita/inferior 2,5 cm (col. única)
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(3.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(3.0)


# ---------------------------------------------------------------- helpers
def body(text, justify=True, italic=False, size=12, space_after=6, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = FONT
    return p


def h1(t):
    return doc.add_heading(t, level=1)


def h2(t):
    return doc.add_heading(t, level=2)


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.font.name = FONT
    r2 = p.add_run(text); r2.font.name = FONT
    return p


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell(cell, text, bold=False, size=10, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = FONT


def table(headers, rows, widths=None, caption=None, font=10):
    if caption:  # SBC: legenda da tabela ACIMA
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after = Pt(3)
        r = cap.add_run(caption); r.font.size = Pt(10); r.font.name = FONT; r.bold = True
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        _set_cell(c, htext, bold=True, size=font); _shade(c, HDR_FILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            _set_cell(cells[i], str(val), size=font)
            if ri % 2 == 1:
                _shade(cells[i], ALT_FILL)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(path, caption, width=15.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(os.path.join(OUT, path), width=Cm(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption); r.font.size = Pt(10); r.font.name = FONT; r.bold = True


# ================================================================ TÍTULO SBC
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Cervix Visual AI: Desenvolvimento de um Pipeline de "
                  "Aprendizado Profundo para Classificação de Lesões Cervicais "
                  "em Imagens de Inspeção Visual com Ácido Acético")
r.bold = True; r.font.size = Pt(16); r.font.name = FONT
title.paragraph_format.space_after = Pt(18)

# autores
au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = au.add_run("[Nome do(a) Autor(a)]¹"); r.font.size = Pt(12); r.font.name = FONT
au.paragraph_format.space_after = Pt(12)

af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
af.paragraph_format.space_after = Pt(2)
r = af.add_run("¹[Curso / Departamento] – [Instituição de Ensino] (SIGLA)")
r.font.size = Pt(12); r.font.name = FONT
af2 = doc.add_paragraph(); af2.alignment = WD_ALIGN_PARAGRAPH.CENTER
af2.paragraph_format.space_after = Pt(2)
r = af2.add_run("[Endereço] – [Cidade] – [UF] – Brasil"); r.font.size = Pt(12); r.font.name = FONT
em = doc.add_paragraph(); em.alignment = WD_ALIGN_PARAGRAPH.CENTER
em.paragraph_format.space_after = Pt(16)
r = em.add_run("ueldomiguel@gmail.com"); r.font.size = Pt(12); r.font.name = FONT


def abstract_block(head, text):
    hp = doc.add_paragraph(); hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(4)
    r = hp.add_run(head); r.bold = True; r.font.size = Pt(12); r.font.name = FONT
    bp = doc.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bp.paragraph_format.left_indent = Cm(0.8)
    bp.paragraph_format.right_indent = Cm(0.8)
    bp.paragraph_format.space_after = Pt(12)
    r = bp.add_run(text); r.italic = True; r.font.size = Pt(12); r.font.name = FONT


# SBC: Abstract (inglês) primeiro, depois Resumo (português)
abstract_block(
    "Abstract.",
    "This technical report describes the development of Cervix Visual AI, a "
    "single-file Python deep-learning pipeline for binary classification of "
    "post-acetic-acid cervical images from the IARC Cervical Image Bank into "
    "negative_or_low_grade and high_grade_or_cancer. We report the full "
    "development process — proposal, implementation and evaluation — covering "
    "automated data preparation, patient-level stratified splitting, aggressive "
    "data augmentation, an EfficientNet-B3 backbone fine-tuned in two stages, a "
    "class-weighted Focal loss, MixUp regularization, cosine-annealing scheduling, "
    "four-view test-time augmentation, temperature-scaling calibration, "
    "sensitivity-constrained threshold selection, case-level aggregation, "
    "bootstrap confidence intervals and Grad-CAM explainability. The software "
    "architecture and its object-oriented design are documented with UML "
    "diagrams. Results are framed as an exploratory research artifact without "
    "established clinical validity.")
abstract_block(
    "Resumo.",
    "Este relatório técnico descreve o desenvolvimento do Cervix Visual AI, um "
    "pipeline de aprendizado profundo em arquivo único (Python) para "
    "classificação binária de imagens cervicais pós-ácido acético do IARC "
    "Cervical Image Bank em negative_or_low_grade e high_grade_or_cancer. "
    "Relata-se o processo completo — proposta, desenvolvimento e avaliação — "
    "abrangendo preparação automatizada dos dados, divisão estratificada por "
    "paciente, aumento agressivo de dados, uma rede EfficientNet-B3 ajustada em "
    "duas fases, perda Focal ponderada por classe, regularização por MixUp, "
    "escalonamento de taxa por cosseno, aumento em tempo de teste de quatro "
    "vistas, calibração por temperatura, seleção de limiar com restrição de "
    "sensibilidade, agregação por caso, intervalos de confiança via bootstrap e "
    "explicabilidade por Grad-CAM. A arquitetura de software e seu projeto "
    "orientado a objetos são documentados com diagramas UML. Os resultados são "
    "apresentados como artefato de pesquisa exploratória, sem validade clínica "
    "estabelecida.")

# ================================================================ 1. INTRODUÇÃO
h1("1. Introdução")
body("O câncer do colo do útero é o quarto tipo de câncer mais comum entre mulheres em todo o "
     "mundo, com aproximadamente 570 mil novos casos e cerca de 342 mil óbitos estimados em "
     "2020, concentrando-se de forma desproporcional em países de baixa e média renda [12]. Em "
     "contextos de recursos limitados, a inspeção visual com ácido acético (IVA) continua sendo "
     "amplamente utilizada como estratégia de rastreamento de baixo custo, apesar de sua "
     "reconhecida dependência do examinador e de sua variabilidade interobservador [1].")
body("Técnicas de visão computacional e aprendizado profundo têm sido propostas como apoio à "
     "decisão para padronizar a interpretação de exames de IVA. Um estudo de referência do "
     "National Cancer Institute demonstrou que a avaliação visual automatizada baseada em "
     "aprendizado profundo identificou casos de pré-câncer/câncer com AUC de 0,91, superando a "
     "interpretação humana (AUC 0,69) e a citologia convencional (AUC 0,71) [3]. A International "
     "Agency for Research on Cancer (IARC) constituiu o IARC Cervical Image Bank para suprir a "
     "carência de bancos de imagens padronizados para o desenvolvimento de tais algoritmos [2].")
body("Este relatório técnico descreve o desenvolvimento do Cervix Visual AI, um pipeline de "
     "visão computacional implementado em um único arquivo Python para a classificação binária "
     "de imagens “After acetic acid” em negative_or_low_grade e high_grade_or_cancer. O texto "
     "está organizado de modo a percorrer as três etapas do trabalho: a proposta (Seção 3), "
     "com objetivos, requisitos e metodologia de desenvolvimento; o desenvolvimento (Seção 4), "
     "com a arquitetura de software e a descrição técnica de cada componente; e a avaliação "
     "(Seção 5), com o protocolo de métricas, os artefatos produzidos e a evolução do pipeline. "
     "A Seção 2 apresenta trabalhos relacionados, a Seção 6 discute ética, limitações e "
     "trabalhos futuros, e a Seção 7 conclui. Cabe ressaltar que, conforme declarado na própria "
     "documentação do código, trata-se de um artefato experimental cujos resultados não devem "
     "ser utilizados para decisão clínica.")

# ================================================================ 2. TRAB. RELACIONADOS
h1("2. Trabalhos Relacionados")
body("O estudo de Hu et al. [3] estabeleceu empiricamente que modelos de aprendizado profundo "
     "podem superar a avaliação visual humana não assistida na triagem cervical, reportando AUC "
     "de 0,91 (IC 95% 0,89–0,93). Especificamente sobre o IARC Cervical Image Bank, um trabalho "
     "recente combinou características de variantes ResNet, reduzidas por LDA e classificadas por "
     "SVM, reportando especificidade, sensibilidade e acurácia de 97% a 99% na tarefa "
     "normal/anormal [5]. Em abordagem multimodal, o sistema DeepCerviCancer combina modelos de "
     "colposcopia e citologia, atingindo valores elevados de F1 e precisão quando ambas as "
     "modalidades estão disponíveis [4]. Para cenários de baixos recursos, um pipeline leve com "
     "EfficientDet-Lite e MobileNet, executado em dispositivo Android, reportou acurácia de "
     "92,31%, sensibilidade de 98,24% e especificidade de 88,37% [6].")
body("O Cervix Visual AI diferencia-se pela combinação de rigor metodológico no tratamento dos "
     "dados (divisão por paciente com verificação de vazamento e agregação por caso) com um "
     "conjunto moderno de técnicas de regularização, calibração e interpretabilidade — perda "
     "Focal, MixUp, escalonamento por cosseno, TTA de quatro vistas, calibração por temperatura, "
     "limiar com restrição de sensibilidade, intervalos de confiança via bootstrap e mapas "
     "Grad-CAM — frequentemente subespecificados em pipelines comparáveis.")

# ================================================================ 3. PROPOSTA
h1("3. Proposta")
h2("3.1 Objetivos")
body("O objetivo geral do trabalho foi desenvolver um pipeline reprodutível, de ponta a ponta, "
     "para a classificação binária de imagens de IVA do IARC Cervical Image Bank, com ênfase em "
     "robustez estatística e interpretabilidade. Como objetivos específicos, definiram-se: (i) "
     "automatizar a preparação dos dados a partir do arquivo distribuído pela IARC; (ii) evitar "
     "vazamento de dados por meio de divisão e avaliação no nível do paciente; (iii) treinar uma "
     "rede convolucional moderna por transferência de aprendizado; (iv) calibrar tanto as "
     "probabilidades quanto o limiar de decisão segundo prioridades clínicas; (v) quantificar a "
     "incerteza das métricas; e (vi) gerar explicações visuais das predições.")
h2("3.2 Requisitos e decisões de projeto")
body("A partir dos objetivos, estabeleceram-se requisitos que orientaram as decisões de projeto, "
     "resumidos na Tabela 1.")
table(
    ["Requisito", "Decisão de projeto"],
    [
        ["Reprodutibilidade e simplicidade de distribuição",
         "Implementação em arquivo único Python, com semente fixa e CLI por subcomandos."],
        ["Ausência de vazamento de dados",
         "Divisão estratificada por patient_id e agregação de predições por caso."],
        ["Desempenho com poucos dados",
         "Transferência de aprendizado, aumento agressivo de dados, MixUp e SWA-like scheduling."],
        ["Sensibilidade clínica (evitar falsos negativos)",
         "Perda Focal com pos_weight e limiar que maximiza F1 sob Sens ≥ 0,80."],
        ["Confiabilidade das probabilidades",
         "Calibração por escalonamento de temperatura (temperature scaling)."],
        ["Quantificação de incerteza",
         "Intervalos de confiança de 95% via bootstrap por caso."],
        ["Interpretabilidade",
         "Mapas Grad-CAM por classe e por quadrante da matriz de confusão."],
    ],
    widths=[5.5, 9.5],
    caption="Tabela 1. Requisitos do projeto e respectivas decisões de projeto.",
)
h2("3.3 Metodologia de desenvolvimento")
body("O desenvolvimento seguiu um processo iterativo e incremental. Uma primeira versão, "
     "baseada em ResNet-34 e com função de perda BCE com suavização de rótulos, serviu de prova "
     "de conceito e expôs limitações de implementação (descritas na Seção 5.3). A versão atual, "
     "documentada neste relatório, evoluiu para um backbone EfficientNet-B3, incorporou perda "
     "Focal, MixUp, escalonamento por cosseno com reinícios, TTA de quatro vistas, calibração "
     "por temperatura, limiar com restrição de sensibilidade e explicabilidade por Grad-CAM, "
     "além de corrigir defeitos identificados na iteração anterior. Cada incremento foi validado "
     "pelo protocolo de avaliação da Seção 5, comparando-se métricas por caso entre versões.")

# ================================================================ 4. DESENVOLVIMENTO
h1("4. Desenvolvimento")
h2("4.1 Arquitetura de software e visão de processo")
body("O Cervix Visual AI é estruturado como uma arquitetura em camadas, com responsabilidades "
     "bem delimitadas e fluxo de controle descendente. A Figura 1 apresenta a visão de processo "
     "— a sequência de transformações aplicadas aos dados, da leitura do arquivo da IARC até a "
     "geração de métricas e mapas de explicabilidade.")
figure("fig1_pipeline.png",
       "Figura 1. Visão de processo do pipeline Cervix Visual AI, da ingestão dos dados à "
       "avaliação por caso, calibração e interpretabilidade.", width=15.5)
body("A Figura 2 apresenta a visão estrutural complementar: a organização do software em "
     "camadas (Interface, Orquestração, Dados, Modelo, Treino e Otimização, Inferência e "
     "Calibração, Avaliação e Métricas, Explicabilidade e Saída) e sua dependência das "
     "bibliotecas externas. Cada camada depende apenas de serviços expostos pelas camadas "
     "inferiores, o que reduz o acoplamento e facilita a substituição de componentes.")
figure("fig_arch.png",
       "Figura 2. Arquitetura de software em camadas; à direita, as bibliotecas externas "
       "utilizadas (PyTorch/torchvision, scikit-learn, pandas/NumPy, matplotlib/Plotly, "
       "OpenCV/PIL).", width=15.5)

h2("4.2 Dados: preparação e divisão por paciente")
body("O pipeline pressupõe o acesso a um arquivo ZIP do IARC Cervical Image Bank, contendo as "
     "planilhas “Cases - Images.xlsx” e “Cases Meta data.xlsx” e os arquivos de imagem [2]. A "
     "função prepare_iarc_dataset() restringe-se às imagens do tipo “After acetic acid”, extrai "
     "cada imagem e gera dois manifestos (manifest.csv, por imagem; case_manifest.csv, por "
     "caso). O rótulo binário não é fornecido diretamente: é derivado por mineração de texto do "
     "campo “Provisional diagnosis” — normalize_text() limpa o texto e derive_target() atribui "
     "high_grade_or_cancer caso haja termos como “hsil”, “cin2/3”, “carcinoma” ou “cancer”, e "
     "negative_or_low_grade caso contrário. Trata-se de rotulagem fraca (weak labeling), cuja "
     "limitação (ausência de tratamento de negação) é discutida na Seção 6.")
body("A divisão dos dados é estratificada e realizada no nível do paciente (Figura 3): 70% dos "
     "casos para treino e os 30% restantes divididos igualmente entre validação e teste "
     "(≈70/15/15). Como um caso pode contribuir com múltiplas imagens, dividir por imagem "
     "configuraria vazamento de dados; load_manifest() verifica que cada paciente pertence a "
     "uma única divisão, e a avaliação é sempre agregada por caso.")
figure("fig4_split.png",
       "Figura 3. Divisão estratificada no nível do paciente, evitando vazamento de dados entre "
       "as divisões.", width=15.0)

h2("4.3 Modelo, aumento de dados e treinamento")
body("A classe CervixImageDataset carrega cada imagem, aplica as transformações e valida o "
     "tensor (NaN/Inf). O pipeline de treino adota um esquema de aumento de dados agressivo "
     "(Tabela 2), voltado a mitigar o sobreajuste em um conjunto médico pequeno; o pipeline de "
     "avaliação é determinístico (redimensionamento, tensor e normalização ImageNet).")
table(
    ["Operação de aumento (treino)", "Parâmetros"],
    [
        ["Resize + RandomCrop", "sobredimensiona +15% e recorta para image_size"],
        ["RandomHorizontalFlip / VerticalFlip", "p = 0,50 / p = 0,30"],
        ["RandomRotation", "± 30°"],
        ["RandomPerspective", "distortion_scale = 0,35; p = 0,50"],
        ["RandomAffine", "rot 20°, translação 12%, escala 0,80–1,20, shear 12"],
        ["ColorJitter", "brilho 0,45; contraste 0,40; saturação 0,35; matiz 0,12"],
        ["Sharpness / Autocontrast / GaussianBlur / Grayscale", "p = 0,40 / 0,30 / sempre / 0,05"],
        ["RandomErasing (×2)", "p = 0,35 e p = 0,25 (escalas distintas)"],
    ],
    widths=[8.5, 6.5],
    caption="Tabela 2. Operações de aumento de dados aplicadas ao conjunto de treino.",
)
body("A função build_model() instancia, por padrão, uma EfficientNet-B3 pré-treinada na "
     "ImageNet [13], escolhida pelo bom equilíbrio entre acurácia e custo proporcionado pelo "
     "escalonamento composto. A Tabela 3 lista as quatro arquiteturas efetivamente suportadas.")
table(
    ["Arquitetura", "Observações"],
    [
        ["efficientnet_b3", "Padrão; escalonamento composto; equilíbrio acurácia/custo [13]"],
        ["efficientnet_b4", "Maior capacidade e campo receptivo [13]"],
        ["resnet34", "Linha de base com conexões residuais [9]"],
        ["convnext_tiny", "ConvNet moderna; boa calibração [14]"],
    ],
    widths=[4.5, 10.5],
    caption="Tabela 3. Arquiteturas de backbone suportadas por build_model().",
)
body("Independentemente do backbone, a camada de classificação é substituída por uma cabeça de "
     "duas camadas (Figura 4): Dropout(0,40) → Linear(256) → ReLU → Dropout(0,20) → Linear(1). "
     "A transferência de aprendizado ocorre em duas fases (Figura 5): nas primeiras cinco "
     "épocas o backbone é congelado e apenas a cabeça é treinada; em seguida, a rede inteira é "
     "ajustada, com taxa de aprendizado do backbone igual a 0,1× a da cabeça.")
figure("fig2_head.png",
       "Figura 4. Cabeça de classificação de duas camadas sobre o vetor de características do "
       "backbone.", width=15.5)
figure("fig3_finetune.png",
       "Figura 5. Transferência de aprendizado em duas fases (congelamento inicial seguido de "
       "ajuste fino completo).", width=15.0)
body("O desbalanceamento é tratado pelo pos_weight da perda, calculado por caso e multiplicado "
     "por um fator (1,5). Duas perdas estão disponíveis (Tabela 4): a Focal Loss (γ = 2,0), "
     "padrão, que concentra o aprendizado nos exemplos difíceis [15]; e a BCE com suavização de "
     "rótulos. O otimizador é o AdamW [10] com taxas diferenciadas, e o agendamento usa "
     "CosineAnnealingWarmRestarts (T_0 = 10, T_mult = 2). O laço de treino aplica MixUp [18] e "
     "recorte de norma de gradiente, e calcula métricas de treino e validação por época. A "
     "Tabela 5 resume os hiperparâmetros padrão.")
table(
    ["Função de perda", "Descrição"],
    [
        ["FocalLoss (padrão)", "BCE modulada por (1 − p_t)^γ, γ = 2,0; foca exemplos difíceis [15]"],
        ["LabelSmoothingBCE", "BCE com suavização de rótulos de 0,05; melhora a calibração"],
    ],
    widths=[4.5, 10.5],
    caption="Tabela 4. Funções de perda disponíveis (ambas com pos_weight).",
)
table(
    ["Hiperparâmetro", "Valor padrão"],
    [
        ["architecture / pretrained", "efficientnet_b3 / True (ImageNet)"],
        ["image_size / batch_size", "384 × 384 / 16 (8 na CLI)"],
        ["epochs / freeze_backbone_epochs / patience", "30 / 5 / 20"],
        ["learning_rate (cabeça) / backbone", "3 × 10⁻⁴ / 0,1×"],
        ["weight_decay", "1 × 10⁻²"],
        ["loss_type / pos_weight_multiplier", "focal (γ = 2,0) / 1,5"],
        ["mixup_alpha / dropout", "0,4 / 0,40 e 0,20"],
        ["scheduler", "CosineAnnealingWarmRestarts (T_0=10, T_mult=2)"],
        ["temperature_scaling / bootstrap / seed", "True / 500 / 42"],
    ],
    widths=[7.0, 8.0],
    caption="Tabela 5. Hiperparâmetros padrão de treinamento.",
)

h2("4.4 Inferência: TTA, calibração e limiar")
body("Na inferência, predict_loader() aplica aumento em tempo de teste com quatro vistas "
     "geométricas (original, espelhamentos horizontal, vertical e duplo), cujas probabilidades "
     "são calculadas em média (Figura 6). As predições por imagem são agregadas por paciente "
     "(média), e optimize_threshold() seleciona o limiar que maximiza o F1 entre aqueles que "
     "garantem sensibilidade ≥ 0,80 — refletindo a prioridade clínica de minimizar falsos "
     "negativos. Após o treino, _calibrate_temperature() ajusta um escalar T por L-BFGS [16], "
     "suavizando probabilidades super-confiantes sem alterar o ranqueamento; T é armazenado no "
     "checkpoint e aplicado em toda inferência.")
figure("fig5_tta.png",
       "Figura 6. Inferência: agregação de quatro vistas de TTA, calibração por temperatura e "
       "decisão por limiar com restrição de sensibilidade.", width=15.5)

h2("4.5 Interpretabilidade: Grad-CAM")
body("A explicabilidade visual é fornecida por Grad-CAM [17]: a classe GradCAM registra ganchos "
     "sobre a última camada convolucional do backbone e combina ativações ponderadas pelos "
     "gradientes da classe-alvo para produzir mapas de calor. A função "
     "generate_gradcam_exemplars() seleciona exemplares de cada quadrante da matriz de confusão "
     "(TP/TN/FP/FN), úteis para auditoria de erros, e generate_gradcam_dataset() estende a "
     "geração a todo o conjunto.")

h2("4.6 Organização orientada a objetos e modelagem UML")
body("Quatro classes concentram as responsabilidades de domínio (Tabela 6). A Figura 7 "
     "apresenta o diagrama de classes UML, e a Figura 8, o diagrama de dependências.")
table(
    ["Classe", "Superclasse", "Responsabilidade"],
    [
        ["CervixImageDataset", "torch.utils.data.Dataset",
         "Carrega imagens, aplica transformações e valida o tensor (NaN/Inf)."],
        ["FocalLoss", "torch.nn.Module",
         "Perda Focal com pos_weight; foca exemplos difíceis."],
        ["LabelSmoothingBCE", "torch.nn.Module",
         "BCE com suavização de rótulos e pos_weight."],
        ["GradCAM", "— (autônoma)",
         "Registra ganchos e gera o mapa de ativação de classe."],
    ],
    widths=[3.6, 4.0, 7.4],
    caption="Tabela 6. Classes do Cervix Visual AI, superclasses e responsabilidades.",
)
body("O diagrama de classes (Figura 7) evidencia três generalizações: FocalLoss e "
     "LabelSmoothingBCE especializam torch.nn.Module, sobrescrevendo forward(); e "
     "CervixImageDataset especializa torch.utils.data.Dataset, implementando __len__() e "
     "__getitem__(). A classe GradCAM não herda de nenhuma superclasse, mas mantém uma "
     "associação com um objeto Module (o modelo a ser explicado) — composição em vez de "
     "herança. Manifestam-se assim os três pilares da orientação a objetos: encapsulamento (o "
     "atributo privado _handles e os métodos _register_hooks()/remove_hooks() de GradCAM); "
     "herança (perdas e dataset reutilizando a infraestrutura do PyTorch); e polimorfismo (a "
     "variável criterion referenciando FocalLoss ou LabelSmoothingBCE, ambas chamadas via "
     "forward()).")
figure("fig6_uml_classes.png",
       "Figura 7. Diagrama de classes UML: classes de domínio, superclasses do PyTorch e a "
       "associação entre GradCAM e o modelo.", width=15.5)
body("A Figura 8 mostra como a função orquestradora train_model() depende de build_model() (que "
     "instancia o modelo), de uma das funções de perda e de create_dataloaders() (que encapsula "
     "objetos CervixImageDataset). Essa separação favorece a manutenibilidade: novas perdas, "
     "backbones ou técnicas de interpretabilidade podem ser adicionados sem alterar o laço de "
     "treinamento.")
figure("fig7_uml_collab.png",
       "Figura 8. Diagrama UML de dependências: train_model() orquestra fábricas, critério de "
       "perda, conjunto de dados e explicabilidade.", width=15.5)

h2("4.7 Interface de linha de comando")
body("O script expõe seis subcomandos (Tabela 7); sem argumentos, assume train.")
table(
    ["Comando", "Função"],
    [
        ["prepare", "Extrai o ZIP da IARC e gera os manifestos."],
        ["demo", "prepare seguido de treino com hiperparâmetros fixos."],
        ["train", "Treino configurável (arquitetura, perda, limiar, temperatura etc.)."],
        ["evaluate", "Reexecuta a avaliação (com TTA e temperatura) sobre um checkpoint."],
        ["predict", "Inferência em imagem única, com probabilidade calibrada e rótulo."],
        ["show", "Exibe imagens ou mapas Grad-CAM no terminal."],
    ],
    widths=[2.6, 12.4],
    caption="Tabela 7. Subcomandos da interface de linha de comando.",
)

# ================================================================ 5. AVALIAÇÃO
h1("5. Avaliação")
h2("5.1 Protocolo e métricas")
body("A avaliação é reportada no nível do caso. A função binary_metrics() calcula acurácia, "
     "acurácia balanceada, sensibilidade, especificidade, precisão (PPV), valor preditivo "
     "negativo, F1 e escore de Brier; quando ambas as classes estão presentes, calcula também "
     "AUC-ROC e precisão média. A seleção do melhor modelo a cada época usa um escore composto "
     "(0,5 × AUC-ROC + 0,3 × F1 + 0,2 × acurácia balanceada), equilibrando discriminação e "
     "desempenho operacional. A função bootstrap_confidence_intervals() realiza 500 "
     "reamostragens com reposição por caso, reportando os percentis 2,5 e 97,5 de cada métrica "
     "como intervalo de confiança aproximado de 95% — adequado às amostras reduzidas de estudos "
     "de centro único.")
h2("5.2 Artefatos produzidos")
body("Uma execução completa do subcomando train produz o conjunto de artefatos da Tabela 8, "
     "incluindo o checkpoint autossuficiente (com limiar e temperatura), os relatórios em JSON, "
     "as predições por caso em CSV, os gráficos e os painéis Grad-CAM.")
table(
    ["Artefato", "Conteúdo"],
    [
        ["best_model.pt", "Pesos, arquitetura, dropout, image_size, limiar e temperatura."],
        ["training_history.json", "Métricas de treino/validação e limiar por época."],
        ["training_summary.json", "Resumo: dispositivo, tempo, peso positivo, melhor época."],
        ["evaluation_metrics.json", "Métricas por divisão (imagem e caso) com IC bootstrap."],
        ["{split}_case_predictions.csv", "Probabilidade e predição por caso, por divisão."],
        ["Gráficos (PNG/HTML)", "Perda, ROC, precisão-revocação, matriz de confusão, painel."],
        ["Painéis Grad-CAM", "Exemplares TP/TN/FP/FN e mapas do conjunto completo."],
    ],
    widths=[5.0, 10.0],
    caption="Tabela 8. Artefatos produzidos por uma execução completa de train.",
)
h2("5.3 Evolução do pipeline")
body("A Tabela 9 sintetiza a evolução entre a versão inicial (ResNet-34) e a versão atual "
     "(EfficientNet-B3), evidenciando tanto correções de defeitos quanto a incorporação de novas "
     "técnicas — o resultado direto da metodologia iterativa descrita na Seção 3.3.")
table(
    ["Aspecto", "Versão inicial (ResNet-34)", "Versão atual (EfficientNet-B3)"],
    [
        ["Backbone", "Apenas ResNet-34 (EfficientNet não construído)",
         "EfficientNet-B3/B4, ResNet-34, ConvNeXt-Tiny"],
        ["Cabeça", "Dropout → Linear(1)", "Dropout → Linear(256) → ReLU → Dropout → Linear(1)"],
        ["Perda", "BCE com suavização", "Focal Loss (padrão) + BCE com suavização"],
        ["Scheduler", "ReduceLROnPlateau", "CosineAnnealingWarmRestarts"],
        ["TTA", "2 vistas", "4 vistas (orig, H, V, HV)"],
        ["Calibração", "Ausente", "Escalonamento de temperatura (L-BFGS)"],
        ["Limiar", "Maximiza F1", "Maximiza F1 com Sens ≥ 0,80"],
        ["Early stopping", "Inoperante (bug)", "Corrigido e funcional"],
        ["Interpretabilidade", "Ausente", "Grad-CAM (exemplares e conjunto)"],
        ["Inferência única", "Retornava logit bruto", "Probabilidade calibrada + rótulo"],
    ],
    widths=[3.0, 6.0, 6.0],
    font=9,
    caption="Tabela 9. Evolução do pipeline entre as iterações de desenvolvimento.",
)
h2("5.4 Resultados esperados e contextualização")
body("Este relatório documenta o projeto e a implementação, e não relata o resultado de uma "
     "execução específica, uma vez que a execução completa depende do acesso ao arquivo "
     "licenciado da IARC, não reproduzido neste trabalho [2]. A título de contextualização — e "
     "não de previsão de desempenho —, trabalhos correlatos sobre o banco da IARC relataram "
     "métricas entre 97% e 99% [5] e, em soluções móveis, acurácia ≈92% [6]. Tais valores "
     "referem-se a conjuntos e definições de rótulo distintos e não devem ser interpretados como "
     "estimativa do desempenho deste pipeline. Recomenda-se que qualquer relato priorize as "
     "métricas no nível do caso, com intervalos de confiança de 95%, o limiar calibrado, a "
     "temperatura e a distribuição de classes por divisão.")

# ================================================================ 6. DISCUSSÃO
h1("6. Discussão: Ética, Limitações e Trabalhos Futuros")
body("O próprio código declara que os resultados são experimentais e não devem ser usados para "
     "decisão clínica — afirmação que este relatório reitera. O acesso às imagens da IARC está "
     "sujeito aos procedimentos da instituição, e o uso de imagens médicas pressupõe aprovação "
     "ética. Há risco de viés algorítmico e de generalização limitada, dada a diversidade de "
     "centros e dispositivos; qualquer aplicação prática exigiria validação clínica prospectiva "
     "[8]. Ferramentas dessa natureza enquadram-se melhor como apoio à decisão do que como "
     "diagnóstico autônomo, e os mapas Grad-CAM contribuem para a transparência.")
body("Permanecem limitações que deveriam ser tratadas antes de qualquer relato empírico:")
bullet("a rotulagem por mineração de texto não trata negação, podendo classificar "
       "incorretamente textos que descrevam a ausência da condição; cotejar com o campo "
       "Histopathology seria uma alternativa.", bold_lead="Rotulagem sem tratamento de negação: ")
bullet("o caminho padrão do ZIP da IARC é específico do ambiente de desenvolvimento, reduzindo "
       "a portabilidade.", bold_lead="Caminho padrão específico: ")
bullet("o paralelismo de múltiplas GPUs está desativado nesta versão.",
       bold_lead="Paralelismo de GPU limitado: ")
bullet("a divisão external_test não é populada automaticamente, embora suportada.",
       bold_lead="Validação externa não automatizada: ")
body("Como trabalhos futuros, delineiam-se: curadoria de rótulos com revisão por especialistas; "
     "validação externa em coorte independente; pré-treinamento de domínio e ensembles por "
     "validação cruzada; fusão multimodal incorporando HPV, escore de Swede e zona de "
     "transformação [4]; e, seguindo [8], validação clínica prospectiva.")

# ================================================================ 7. CONCLUSÃO
h1("7. Conclusão")
body("Este relatório técnico descreveu o desenvolvimento completo do Cervix Visual AI, "
     "percorrendo a proposta (objetivos, requisitos e metodologia iterativa), o desenvolvimento "
     "(arquitetura de software, dados, modelo EfficientNet-B3, treinamento, inferência "
     "calibrada, interpretabilidade e projeto orientado a objetos documentado em UML) e a "
     "avaliação (protocolo por caso, artefatos, evolução entre versões e contextualização de "
     "resultados). O pipeline combina disciplina no tratamento de dados — divisão e agregação "
     "por paciente com quantificação de incerteza — com técnicas modernas de regularização, "
     "calibração e explicabilidade, corrigindo defeitos da iteração inicial. Enquadrado "
     "honestamente como artefato de pesquisa experimental, oferece um modelo reprodutível e "
     "interpretável, além de um ponto de partida concreto para refinamentos — curadoria de "
     "rótulos, validação externa, pré-treinamento de domínio, fusão multimodal e avaliação "
     "clínica prospectiva — necessários para que um protótipo evolua, com segurança para as "
     "pacientes, em direção a uma ferramenta de apoio ao rastreamento do câncer do colo do "
     "útero [1].")

# ================================================================ REFERÊNCIAS
h1("Referências")
refs = [
    "WORLD HEALTH ORGANIZATION. Global strategy to accelerate the elimination of cervical "
    "cancer as a public health problem. Geneva: WHO, 2020.",
    "INTERNATIONAL AGENCY FOR RESEARCH ON CANCER (IARC). IARC Cervical Cancer Image Bank. "
    "Disponível em: https://screening.iarc.fr/cervicalimagebank.php. Acesso em: jun. 2026.",
    "HU, L. et al. An Observational Study of Deep Learning and Automated Evaluation of Cervical "
    "Images for Cancer Screening. Journal of the National Cancer Institute, 2019.",
    "KALBHOR, M. et al. DeepCerviCancer — Deep Learning-Based Cervical Image Classification "
    "using Colposcopy and Cytology Images. EAI Endorsed Trans. on Pervasive Health and "
    "Technology, 2023.",
    "Deep Learning Descriptor Hybridization with Feature Reduction for Accurate Cervical Cancer "
    "Colposcopy Image Classification. arXiv:2405.01600, 2024.",
    "Automated Cervical Cancer Detection through Visual Inspection with Acetic Acid in "
    "Resource-Poor Settings with Lightweight Deep Learning Models Deployed on an Android Device. "
    "arXiv:2508.13253, 2025.",
    "XUE, Z. et al. Image Quality Classification for Automated Visual Evaluation of Cervical "
    "Precancer. In: MILLanD 2022, LNCS v. 13559. Cham: Springer, 2022.",
    "Development and Clinical Validation of a VIA-Artificial Intelligence Tool in "
    "Screen-and-Treat Visual Screening for Cervical Cancer in South India: A Pilot Study. JCO "
    "Global Oncology, 2024.",
    "HE, K.; ZHANG, X.; REN, S.; SUN, J. Deep Residual Learning for Image Recognition. In: "
    "CVPR, 2016.",
    "LOSHCHILOV, I.; HUTTER, F. Decoupled Weight Decay Regularization. In: ICLR, 2019.",
    "DENG, J. et al. ImageNet: A Large-Scale Hierarchical Image Database. In: CVPR, 2009.",
    "WORLD HEALTH ORGANIZATION. New recommendations for screening and treatment to prevent "
    "cervical cancer. Geneva: WHO, 2021.",
    "TAN, M.; LE, Q. EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks. "
    "In: ICML, 2019.",
    "LIU, Z. et al. A ConvNet for the 2020s. In: CVPR, 2022.",
    "LIN, T.-Y. et al. Focal Loss for Dense Object Detection. In: ICCV, 2017.",
    "GUO, C. et al. On Calibration of Modern Neural Networks. In: ICML, 2017.",
    "SELVARAJU, R. R. et al. Grad-CAM: Visual Explanations from Deep Networks via "
    "Gradient-based Localization. In: ICCV, 2017.",
    "ZHANG, H. et al. mixup: Beyond Empirical Risk Minimization. In: ICLR, 2018.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"[{i}] {ref}"); r.font.size = Pt(11); r.font.name = FONT

out_path = os.path.join(OUT, "cervix_visual_ai_trabalho_academico.docx")
doc.save(out_path)
print("Documento SBC salvo:", out_path)
print("Parágrafos:", len(doc.paragraphs), "| Tabelas:", len(doc.tables))
